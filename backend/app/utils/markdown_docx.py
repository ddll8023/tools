"""将常用 Markdown 结构渲染为 DOCX。"""

from __future__ import annotations

import base64
import binascii
import re
from dataclasses import dataclass, replace
from io import BytesIO
from pathlib import Path
from typing import Iterable, Sequence
from urllib.parse import unquote, urlsplit

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token


_MAX_DATA_URI_SIZE = 10 * 1024 * 1024
_MAX_IMAGE_WIDTH_INCHES = 6.0


@dataclass(frozen=True)
class _RenderContext:
    """块级渲染上下文。"""

    list_kind: str | None = None
    list_depth: int = 0
    blockquote_depth: int = 0


@dataclass(frozen=True)
class _InlineStyle:
    """行内文本样式。"""

    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False


class MarkdownDocxRenderer:
    """将 Markdown Token 渲染到 python-docx 文档。"""

    def __init__(self, document: DocumentType, base_dir: Path):
        self.document = document
        self.base_dir = base_dir.resolve()
        self.warnings: list[str] = []
        self._warning_set: set[str] = set()
        self._configure_document()

    def render(self, tokens: Sequence[Token]) -> list[str]:
        self._render_blocks(tokens, 0, len(tokens), _RenderContext())
        return self.warnings

    def _configure_document(self) -> None:
        """设置文档页边距和常用样式。"""
        for section in self.document.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        normal = self.document.styles["Normal"]
        normal.font.name = "Aptos"
        normal.font.size = Pt(11)

        for level in range(1, 7):
            style = self.document.styles[f"Heading {level}"]
            style.font.name = "Aptos Display"

        if "Code Block" not in self.document.styles:
            code_style = self.document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
            code_style.base_style = normal
            code_style.font.name = "Consolas"
            code_style.font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.25)
            code_style.paragraph_format.right_indent = Inches(0.25)
            code_style.paragraph_format.space_before = Pt(4)
            code_style.paragraph_format.space_after = Pt(4)

    def _render_blocks(
        self,
        tokens: Sequence[Token],
        start: int,
        end: int,
        context: _RenderContext,
    ) -> None:
        index = start
        while index < end:
            token = tokens[index]

            if token.type in {"paragraph_open", "heading_open"}:
                close_type = "paragraph_close" if token.type == "paragraph_open" else "heading_close"
                close_index = self._find_simple_close(tokens, index, close_type, end)
                inline = tokens[index + 1] if index + 1 < close_index else None
                paragraph = self._add_block_paragraph(token, context)
                if isinstance(inline, Token) and inline.type == "inline":
                    self._render_inline(inline.children or [], paragraph)
                index = close_index + 1
                continue

            if token.type in {"bullet_list_open", "ordered_list_open"}:
                close_index = self._find_list_close(tokens, index, end)
                kind = "ordered" if token.type == "ordered_list_open" else "bullet"
                self._render_list(tokens, index + 1, close_index, context, kind)
                index = close_index + 1
                continue

            if token.type == "blockquote_open":
                close_index = self._find_matching(
                    tokens,
                    index,
                    {"blockquote_open"},
                    {"blockquote_close"},
                    end,
                )
                quote_context = replace(
                    context,
                    blockquote_depth=context.blockquote_depth + 1,
                )
                self._render_blocks(tokens, index + 1, close_index, quote_context)
                index = close_index + 1
                continue

            if token.type == "table_open":
                close_index = self._find_matching(
                    tokens,
                    index,
                    {"table_open"},
                    {"table_close"},
                    end,
                )
                self._render_table(tokens[index + 1 : close_index])
                index = close_index + 1
                continue

            if token.type in {"fence", "code_block"}:
                self._add_code_block(token.content)
                index += 1
                continue

            if token.type == "hr":
                paragraph = self.document.add_paragraph()
                paragraph.add_run("─" * 42)
                index += 1
                continue

            if token.type in {"html_block", "html_inline"}:
                self._warn("已忽略原始 HTML 内容")

            index += 1

    def _add_block_paragraph(
        self,
        token: Token,
        context: _RenderContext,
    ):
        if token.type == "heading_open":
            level = max(1, min(6, int(token.tag[1:]) if token.tag.startswith("h") else 1))
            return self.document.add_paragraph(style=f"Heading {level}")

        if context.list_kind:
            depth = max(1, min(3, context.list_depth))
            style_prefix = "List Number" if context.list_kind == "ordered" else "List Bullet"
            style_name = style_prefix if depth == 1 else f"{style_prefix} {depth}"
            try:
                return self.document.add_paragraph(style=style_name)
            except KeyError:
                return self.document.add_paragraph(style=style_prefix)

        if context.blockquote_depth:
            try:
                return self.document.add_paragraph(style="Intense Quote")
            except KeyError:
                return self.document.add_paragraph(style="Quote")

        return self.document.add_paragraph()

    def _render_list(
        self,
        tokens: Sequence[Token],
        start: int,
        end: int,
        context: _RenderContext,
        kind: str,
    ) -> None:
        index = start
        item_context = replace(
            context,
            list_kind=kind,
            list_depth=context.list_depth + 1,
        )
        while index < end:
            if tokens[index].type != "list_item_open":
                index += 1
                continue
            item_close = self._find_matching(
                tokens,
                index,
                {"list_item_open"},
                {"list_item_close"},
                end,
            )
            self._render_blocks(tokens, index + 1, item_close, item_context)
            index = item_close + 1

    def _render_table(self, tokens: Sequence[Token]) -> None:
        rows: list[tuple[bool, list[list[Token]]]] = []
        current_row: list[list[Token]] | None = None
        current_cell: list[Token] | None = None
        current_header = False
        row_header = False

        for token in tokens:
            if token.type == "tr_open":
                current_row = []
                row_header = False
            elif token.type in {"th_open", "td_open"}:
                if current_row is None:
                    continue
                current_cell = []
                current_header = token.type == "th_open"
                row_header = row_header or current_header
            elif token.type == "inline" and current_cell is not None:
                current_cell.extend(token.children or [])
            elif token.type in {"th_close", "td_close"}:
                if current_row is not None and current_cell is not None:
                    current_row.append(current_cell)
                current_cell = None
            elif token.type == "tr_close":
                if current_row:
                    rows.append((row_header, current_row))
                current_row = None

        if not rows:
            self._warn("表格没有可转换的内容")
            return

        column_count = max(len(cells) for _, cells in rows)
        table = self.document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"
        table.autofit = True

        for row_index, (is_header, cells) in enumerate(rows):
            for column_index in range(column_count):
                cell = table.cell(row_index, column_index)
                cell.text = ""
                paragraph = cell.paragraphs[0]
                if column_index < len(cells):
                    self._render_inline(cells[column_index], paragraph)
                if is_header:
                    for run in paragraph.runs:
                        run.bold = True
                    self._shade_cell(cell, "F2F2F2")

    def _add_code_block(self, content: str) -> None:
        paragraph = self.document.add_paragraph(style="Code Block")
        paragraph.add_run(content.rstrip("\n"))
        self._shade_paragraph(paragraph, "F4F4F4")

    def _render_inline(self, tokens: Iterable[Token], paragraph) -> None:
        style_stack = [_InlineStyle()]
        token_list = list(tokens)
        index = 0

        while index < len(token_list):
            token = token_list[index]

            if token.type == "link_open":
                close_index = self._find_inline_close(token_list, index, "link_close")
                href = token.attrGet("href") or ""
                text = self._inline_plain_text(token_list[index + 1 : close_index])
                if text:
                    self._add_hyperlink(paragraph, href, text)
                index = close_index + 1
                continue

            if token.type in {"strong_open", "em_open", "s_open"}:
                current = style_stack[-1]
                style_stack.append(
                    replace(
                        current,
                        bold=current.bold or token.type == "strong_open",
                        italic=current.italic or token.type == "em_open",
                        strike=current.strike or token.type == "s_open",
                    )
                )
                index += 1
                continue

            if token.type in {"strong_close", "em_close", "s_close"}:
                if len(style_stack) > 1:
                    style_stack.pop()
                index += 1
                continue

            if token.type in {"text", "code_inline"}:
                style = style_stack[-1]
                if token.type == "code_inline":
                    style = replace(style, code=True)
                run = paragraph.add_run(token.content)
                self._apply_run_style(run, style)
                index += 1
                continue

            if token.type in {"softbreak", "hardbreak"}:
                paragraph.add_run().add_break()
                index += 1
                continue

            if token.type == "image":
                self._add_image(paragraph, token.attrGet("src") or "", token.content)
                index += 1
                continue

            if token.type == "html_inline":
                self._warn("已忽略原始 HTML 内容")
                index += 1
                continue

            if token.content:
                run = paragraph.add_run(token.content)
                self._apply_run_style(run, style_stack[-1])
            index += 1

    def _add_image(self, paragraph, source: str, alt_text: str) -> None:
        image = self._resolve_image(source)
        if image is None:
            paragraph.add_run(f"[图片未找到: {alt_text or source}]")
            return

        run = paragraph.add_run()
        try:
            if isinstance(image, BytesIO):
                run.add_picture(image, width=Inches(_MAX_IMAGE_WIDTH_INCHES))
            else:
                run.add_picture(str(image), width=Inches(_MAX_IMAGE_WIDTH_INCHES))
        except Exception:
            self._warn(f"图片无法嵌入: {alt_text or source}")
            paragraph.add_run(f"[图片无法嵌入: {alt_text or source}]")

    def _resolve_image(self, source: str) -> Path | BytesIO | None:
        if not source:
            self._warn("Markdown 图片缺少路径")
            return None

        if source.startswith("data:"):
            return self._decode_data_uri(source)

        parsed = urlsplit(source)
        if parsed.scheme or parsed.netloc:
            self._warn(f"已跳过远程图片: {source}")
            return None

        relative_path = Path(unquote(parsed.path))
        if relative_path.is_absolute() or ".." in relative_path.parts:
            self._warn(f"已跳过越界图片路径: {source}")
            return None

        candidate = (self.base_dir / relative_path).resolve()
        if not candidate.is_relative_to(self.base_dir):
            self._warn(f"已跳过越界图片路径: {source}")
            return None
        if not candidate.is_file():
            self._warn(f"图片文件不存在: {source}")
            return None
        return candidate

    def _decode_data_uri(self, source: str) -> BytesIO | None:
        match = re.match(r"^data:[^;,]+(?:;[^;,]+)*;base64,(.+)$", source, re.IGNORECASE | re.DOTALL)
        if not match:
            self._warn("已跳过不支持的 Data URI 图片")
            return None
        try:
            data = base64.b64decode(match.group(1), validate=True)
        except (ValueError, binascii.Error):
            self._warn("Data URI 图片内容无效")
            return None
        if len(data) > _MAX_DATA_URI_SIZE:
            self._warn("Data URI 图片超过 10MB，已跳过")
            return None
        return BytesIO(data)

    @staticmethod
    def _apply_run_style(run, style: _InlineStyle) -> None:
        run.bold = style.bold
        run.italic = style.italic
        run.font.strike = style.strike
        if style.code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(160, 32, 64)

    @staticmethod
    def _add_hyperlink(paragraph, href: str, text: str) -> None:
        parsed = urlsplit(href)
        if parsed.scheme not in {"http", "https", "mailto"}:
            paragraph.add_run(text)
            return

        relationship_id = paragraph.part.relate_to(
            href,
            RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)

        run = OxmlElement("w:r")
        properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        properties.append(color)
        properties.append(underline)
        run.append(properties)

        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def _inline_plain_text(tokens: Sequence[Token]) -> str:
        parts: list[str] = []
        for token in tokens:
            if token.type in {"text", "code_inline"}:
                parts.append(token.content)
            elif token.type in {"softbreak", "hardbreak"}:
                parts.append("\n")
            elif token.type == "image":
                parts.append(token.content)
        return "".join(parts)

    @staticmethod
    def _find_simple_close(
        tokens: Sequence[Token],
        start: int,
        close_type: str,
        end: int,
    ) -> int:
        for index in range(start + 1, end):
            if tokens[index].type == close_type:
                return index
        return end - 1

    @staticmethod
    def _find_inline_close(tokens: Sequence[Token], start: int, close_type: str) -> int:
        for index in range(start + 1, len(tokens)):
            if tokens[index].type == close_type:
                return index
        return len(tokens)

    @staticmethod
    def _find_matching(
        tokens: Sequence[Token],
        start: int,
        open_types: set[str],
        close_types: set[str],
        end: int,
    ) -> int:
        depth = 0
        for index in range(start, end):
            token_type = tokens[index].type
            if token_type in open_types:
                depth += 1
            elif token_type in close_types:
                depth -= 1
                if depth == 0:
                    return index
        return end

    def _find_list_close(self, tokens: Sequence[Token], start: int, end: int) -> int:
        return self._find_matching(
            tokens,
            start,
            {"bullet_list_open", "ordered_list_open"},
            {"bullet_list_close", "ordered_list_close"},
            end,
        )

    def _warn(self, message: str) -> None:
        if message not in self._warning_set:
            self._warning_set.add(message)
            self.warnings.append(message)

    @staticmethod
    def _shade_paragraph(paragraph, fill: str) -> None:
        properties = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)


def render_markdown_to_docx(
    markdown_content: str,
    base_dir: Path,
    output_path: Path,
) -> list[str]:
    """将 Markdown 文本渲染为 DOCX，并返回非致命转换警告。"""
    parser = MarkdownIt("commonmark", {"html": False})
    parser.enable(["table", "strikethrough"])
    tokens = parser.parse(markdown_content)

    document = Document()
    renderer = MarkdownDocxRenderer(document, base_dir)
    warnings = renderer.render(tokens)
    document.save(str(output_path))
    return warnings
